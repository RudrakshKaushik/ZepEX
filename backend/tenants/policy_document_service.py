import json
import mimetypes
import os
import tempfile
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import List, Optional

from django.conf import settings
from docx import Document
from google import genai
from google.genai import types
from pydantic import BaseModel, Field, ValidationError, field_validator

from .models import (
    CompanyRole,
    Currency,
)
from .prompts import build_policy_document_prompt


from .language_utils import get_company_output_language
# ============================================================
# File configuration
# ============================================================

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
}

MAX_DOCUMENT_SIZE = 25 * 1024 * 1024


# ============================================================
# Categories used by the ZepEx policy engine
# ============================================================

VALID_CATEGORIES = {
    "food",
    "hotel",
    "flight_ticket",
    "train_ticket",
    "car_rental",
    "taxi",
    "fuel",
    "gas",
    "parking",
    "office_supplies",
    "medical",
    "courier",
    "telecom",
    "training",
    "relocation",
    "wfh",
    "miscellaneous",
}


CATEGORY_ALIASES = {
    # Food
    "meal": "food",
    "meals": "food",
    "breakfast": "food",
    "lunch": "food",
    "dinner": "food",
    "snack": "food",
    "snacks": "food",
    "refreshment": "food",
    "refreshments": "food",
    "business meal": "food",
    "client meal": "food",

    # Hotel
    "lodging": "hotel",
    "accommodation": "hotel",
    "hotel stay": "hotel",
    "room charges": "hotel",
    "guest house": "hotel",

    # Flight
    "airfare": "flight_ticket",
    "flight": "flight_ticket",
    "flight ticket": "flight_ticket",
    "air ticket": "flight_ticket",
    "air travel": "flight_ticket",

    # Train
    "rail": "train_ticket",
    "railway": "train_ticket",
    "train": "train_ticket",
    "rail fare": "train_ticket",
    "train fare": "train_ticket",

    # Car rental
    "rental car": "car_rental",
    "vehicle rental": "car_rental",
    "vehicle hire": "car_rental",
    "self drive": "car_rental",

    # Taxi
    "cab": "taxi",
    "uber": "taxi",
    "ola": "taxi",
    "ride": "taxi",
    "local transport": "taxi",

    # Fuel
    "petrol": "fuel",
    "diesel": "fuel",
    "gasoline": "fuel",

    # Gas
    "cng": "gas",
    "lpg": "gas",

    # Parking
    "parking fee": "parking",
    "parking charges": "parking",
    "toll": "parking",
    "toll tax": "parking",

    # Office supplies
    "stationery": "office_supplies",
    "printing": "office_supplies",
    "business supplies": "office_supplies",

    # Medical
    "healthcare": "medical",
    "medicine": "medical",
    "hospital": "medical",
    "pharmacy": "medical",

    # Courier
    "shipping": "courier",
    "postal": "courier",
    "delivery": "courier",

    # Telecom
    "mobile": "telecom",
    "telephone": "telecom",
    "internet": "telecom",
    "broadband": "telecom",
    "data plan": "telecom",

    # Training
    "certification": "training",
    "course": "training",
    "conference": "training",
    "seminar": "training",
    "workshop": "training",

    # Relocation
    "moving": "relocation",
    "transfer expense": "relocation",
    "joining relocation": "relocation",

    # Work from home
    "work from home": "wfh",
    "home office": "wfh",
    "remote work": "wfh",

    # Miscellaneous
    "other": "miscellaneous",
    "general expense": "miscellaneous",
    "uncategorized": "miscellaneous",
}


VALID_LIMIT_PERIODS = {
    "PER_MEAL",
    "PER_DAY",
    "PER_NIGHT",
    "PER_TRIP",
    "PER_MONTH",
    "PER_YEAR",
    "PER_KILOMETRE",
    "PER_MILE",
    "PER_EMPLOYEE",
    "PER_EVENT",
    "PER_INVOICE",
    "ACTUAL_COST",
    "NOT_SPECIFIED",
}


VALID_TRAVEL_SCOPES = {
    "DOMESTIC",
    "INTERNATIONAL",
    "BOTH",
    "NOT_SPECIFIED",
}


# ============================================================
# Gemini structured-output schema
# ============================================================

class DetectedLanguage(BaseModel):
    language: str
    language_code: Optional[str] = None

    confidence: float = Field(
        default=0.0,
        ge=0.0,
        le=1.0,
    )


class ExtractedRoleMatch(BaseModel):
    matched: bool = False
    matched_role_name: Optional[str] = None
    matched_role_id: Optional[str] = None
    match_type: str = "ROLE_NOT_FOUND"

from typing import Literal
class ExtractedPolicyRule(BaseModel):
    role: Optional[str] = None
    role_original_text: Optional[str] = None
    role_match: Optional[ExtractedRoleMatch] = None
    scope: Literal[
        "ALL",
        "ROLE",
    ] = "ALL"
    category: str
    original_category: Optional[str] = None

    max_amount: Optional[str] = None
    currency: Optional[str] = None

    is_unlimited: bool = False
    is_allowed: bool = True

    description: str
    description_original_language: Optional[str] = None

    reason: str
    reason_original_language: Optional[str] = None

    source_text: Optional[str] = None
    source_language: Optional[str] = None
    source_language_code: Optional[str] = None
    translated_source_text: Optional[str] = None

    limit_period: str = "NOT_SPECIFIED"

    conditions: List[str] = Field(default_factory=list)
    conditions_original_language: List[str] = Field(
        default_factory=list
    )

    required_documents: List[str] = Field(
        default_factory=list
    )
    required_documents_original_language: List[str] = Field(
        default_factory=list
    )

    receipt_required: Optional[bool] = None
    approval_required: Optional[bool] = None

    domestic_or_international: str = "NOT_SPECIFIED"

    country_scope: List[str] = Field(default_factory=list)
    department_scope: List[str] = Field(default_factory=list)
    employee_grade_scope: List[str] = Field(default_factory=list)

    travel_class: Optional[str] = None

    distance_rate: Optional[str] = None
    distance_unit: Optional[str] = None

    tax_included: Optional[bool] = None

    confidence: float = Field(
        default=0.0,
        ge=0.0,
        le=1.0,
    )

    translation_confidence: float = Field(
        default=0.0,
        ge=0.0,
        le=1.0,
    )

    review_required: bool = False
    review_reason: Optional[str] = None

    @field_validator("role")
    @classmethod
    def validate_role(cls, value):
        if value is None:
            return None

        cleaned = str(value).strip()

        if cleaned.casefold() in {
            "",
            "none",
            "null",
            "general",
            "default",
            "all employees",
            "all staff",
            "company wide",
            "company-wide",
        }:
            return None

        return cleaned

    @field_validator("category")
    @classmethod
    def validate_category(cls, value):
        if not value:
            return "miscellaneous"

        return str(value).strip().lower()

    @field_validator("currency")
    @classmethod
    def validate_currency(cls, value):
        if not value:
            return None

        return str(value).strip().upper()

    @field_validator("max_amount")
    @classmethod
    def validate_amount(cls, value):
        if value is None:
            return None

        cleaned = (
            str(value)
            .replace(" ", "")
            .replace(",", "")
            .strip()
        )

        if cleaned.casefold() in {
            "",
            "none",
            "null",
            "unlimited",
            "actual",
            "actualcost",
            "actualexpense",
        }:
            return None

        try:
            amount = Decimal(cleaned)
        except InvalidOperation as exc:
            raise ValueError(
                "max_amount must be a numeric value."
            ) from exc

        if amount < 0:
            raise ValueError(
                "max_amount cannot be negative."
            )

        return format(amount, "f")

    @field_validator("distance_rate")
    @classmethod
    def validate_distance_rate(cls, value):
        if value is None:
            return None

        cleaned = (
            str(value)
            .replace(" ", "")
            .replace(",", "")
            .strip()
        )

        if not cleaned:
            return None

        try:
            amount = Decimal(cleaned)
        except InvalidOperation as exc:
            raise ValueError(
                "distance_rate must be numeric."
            ) from exc

        if amount < 0:
            raise ValueError(
                "distance_rate cannot be negative."
            )

        return format(amount, "f")

    @field_validator("limit_period")
    @classmethod
    def validate_limit_period(cls, value):
        normalized = str(
            value or "NOT_SPECIFIED"
        ).strip().upper()

        if normalized not in VALID_LIMIT_PERIODS:
            return "NOT_SPECIFIED"

        return normalized

    @field_validator("domestic_or_international")
    @classmethod
    def validate_travel_scope(cls, value):
        normalized = str(
            value or "NOT_SPECIFIED"
        ).strip().upper()

        if normalized not in VALID_TRAVEL_SCOPES:
            return "NOT_SPECIFIED"

        return normalized


class ExtractedWarning(BaseModel):
    code: str
    message: str
    rule_index: Optional[int] = None
    source_text: Optional[str] = None


class ExtractedConflict(BaseModel):
    role: Optional[str] = None
    category: str
    values: List[str] = Field(default_factory=list)
    message: str
    source_texts: List[str] = Field(default_factory=list)


class ExtractedPolicyDocument(BaseModel):
    document_language: Optional[str] = None
    document_language_code: Optional[str] = None
    contains_multiple_languages: bool = False

    languages_detected: List[DetectedLanguage] = Field(
        default_factory=list
    )

    policy_name: str = "Imported Expense Policy"
    policy_version: Optional[str] = None

    effective_date: Optional[str] = None
    expiry_date: Optional[str] = None

    policy_currency: Optional[str] = None

    document_summary: Optional[str] = None
    document_summary_original_language: Optional[str] = None

    rules: List[ExtractedPolicyRule] = Field(
        default_factory=list
    )

    warnings: List[ExtractedWarning] = Field(
        default_factory=list
    )

    conflicts: List[ExtractedConflict] = Field(
        default_factory=list
    )

    document_level_conditions: List[str] = Field(
        default_factory=list
    )

    document_level_conditions_original_language: List[str] = Field(
        default_factory=list
    )

    @field_validator("policy_currency")
    @classmethod
    def validate_policy_currency(cls, value):
        if not value:
            return None

        return str(value).strip().upper()


# ============================================================
# Configuration helpers
# ============================================================

def _get_gemini_client():
    api_key = getattr(
        settings,
        "GEMINI_API_KEY",
        "",
    )

    if not api_key:
        raise ValueError(
            "GEMINI_API_KEY is not configured."
        )

    return genai.Client(
        api_key=api_key
    )


def _get_policy_model():
    return (
        getattr(
            settings,
            "GEMINI_POLICY_MODEL",
            "",
        )
        or getattr(
            settings,
            "GEMINI_RECEIPT_MODEL",
            "",
        )
        or "gemini-3-flash-preview"
    )


# ============================================================
# File processing
# ============================================================

def _validate_uploaded_file(uploaded_file):
    if not uploaded_file:
        raise ValueError(
            "Policy document is required."
        )

    extension = Path(
        uploaded_file.name or ""
    ).suffix.lower()

    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError(
            "Unsupported policy document. "
            "Allowed formats: PDF, DOCX, TXT, "
            "JPG, JPEG, PNG and WEBP."
        )

    if uploaded_file.size > MAX_DOCUMENT_SIZE:
        raise ValueError(
            "Policy document cannot be larger than 25 MB."
        )

    return extension


def _save_temporary_file(
    uploaded_file,
    extension,
):
    temporary_path = None

    with tempfile.NamedTemporaryFile(
        delete=False,
        suffix=extension,
    ) as temporary_file:
        for chunk in uploaded_file.chunks():
            temporary_file.write(chunk)

        temporary_path = temporary_file.name

    return temporary_path


def _extract_docx_text(file_path):
    document = Document(file_path)

    output = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            output.append(text)

    for table_number, table in enumerate(
        document.tables,
        start=1,
    ):
        output.append(
            f"\nTABLE {table_number}\n"
        )

        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            if any(cells):
                output.append(
                    " | ".join(cells)
                )

    content = "\n".join(output).strip()

    if not content:
        raise ValueError(
            "No readable content was found in the DOCX file."
        )

    return content


def _extract_txt_text(file_path):
    encodings = [
        "utf-8-sig",
        "utf-8",
        "utf-16",
        "latin-1",
    ]

    for encoding in encodings:
        try:
            text = Path(file_path).read_text(
                encoding=encoding
            ).strip()

            if text:
                return text

        except UnicodeDecodeError:
            continue

    raise ValueError(
        "The TXT file could not be decoded."
    )


# ============================================================
# Database reference indexes
# ============================================================

def _get_currency_indexes():
    currencies = Currency.objects.filter(
        is_active=True
    )

    by_code = {}
    by_name = {}
    by_symbol = {}
    by_country = {}

    for currency in currencies:
        code = currency.code.strip().upper()

        by_code[code] = currency

        if currency.name:
            by_name[
                currency.name.strip().casefold()
            ] = currency

        if currency.symbol:
            by_symbol.setdefault(
                currency.symbol.strip(),
                [],
            ).append(currency)

        if currency.country:
            by_country[
                currency.country.strip().casefold()
            ] = currency

    return {
        "by_code": by_code,
        "by_name": by_name,
        "by_symbol": by_symbol,
        "by_country": by_country,
    }


def _get_role_indexes(company):
    roles = CompanyRole.objects.filter(
        company=company,
        is_active=True,
    )

    by_id = {}
    by_name = {}

    for role in roles:
        by_id[str(role.id)] = role
        by_name[
            role.name.strip().casefold()
        ] = role

    return {
        "by_id": by_id,
        "by_name": by_name,
    }


def _get_employee_role(company):
    return CompanyRole.objects.filter(
        company=company,
        name__iexact="Employee",
        is_active=True,
    ).first()


def _get_company_base_currency(company):
    finance_settings = getattr(
        company,
        "finance_settings",
        None,
    )

    if (
        finance_settings
        and finance_settings.base_currency
    ):
        return finance_settings.base_currency

    return None


# ============================================================
# Normalization
# ============================================================

def _normalize_category(
    category,
    original_category=None,
):
    normalized = str(
        category or ""
    ).strip().lower()

    if normalized in VALID_CATEGORIES:
        return normalized, False

    alias_match = CATEGORY_ALIASES.get(
        normalized
    )

    if alias_match:
        return alias_match, False

    original_normalized = str(
        original_category or ""
    ).strip().lower()

    alias_match = CATEGORY_ALIASES.get(
        original_normalized
    )

    if alias_match:
        return alias_match, False

    return "miscellaneous", True


def _resolve_currency(
    currency_value,
    *,
    currency_indexes,
    policy_currency=None,
    base_currency=None,
):
    candidates = [
        currency_value,
        policy_currency,
    ]

    for candidate in candidates:
        if not candidate:
            continue

        raw = str(candidate).strip()
        code = raw.upper()

        if code in currency_indexes["by_code"]:
            return (
                currency_indexes["by_code"][code],
                "CODE",
            )

        normalized = raw.casefold()

        if normalized in currency_indexes["by_name"]:
            return (
                currency_indexes["by_name"][normalized],
                "NAME",
            )

        if normalized in currency_indexes["by_country"]:
            return (
                currency_indexes["by_country"][normalized],
                "COUNTRY",
            )

        symbol_matches = currency_indexes[
            "by_symbol"
        ].get(raw, [])

        if len(symbol_matches) == 1:
            return symbol_matches[0], "SYMBOL"

    if base_currency:
        return (
            base_currency,
            "COMPANY_BASE_CURRENCY",
        )

    return None, None


def _resolve_role(
    *,
    role_value,
    ai_role_match,
    role_indexes,
):
    if not role_value:
        return {
            "role_id": None,
            "role_name": None,
            "matched": True,
            "match_type": "DEFAULT_EMPLOYEE_FALLBACK",
        }

    if (
        ai_role_match
        and ai_role_match.matched_role_id
        and str(ai_role_match.matched_role_id)
        in role_indexes["by_id"]
    ):
        role = role_indexes["by_id"][
            str(ai_role_match.matched_role_id)
        ]

        return {
            "role_id": str(role.id),
            "role_name": role.name,
            "matched": True,
            "match_type": "ROLE_ID",
        }

    normalized_name = str(
        role_value
    ).strip().casefold()

    direct_match = role_indexes[
        "by_name"
    ].get(normalized_name)

    if direct_match:
        return {
            "role_id": str(direct_match.id),
            "role_name": direct_match.name,
            "matched": True,
            "match_type": "EXACT_NAME",
        }

    if (
        ai_role_match
        and ai_role_match.matched_role_name
    ):
        ai_name = (
            ai_role_match.matched_role_name
            .strip()
            .casefold()
        )

        matched_role = role_indexes[
            "by_name"
        ].get(ai_name)

        if matched_role:
            return {
                "role_id": str(matched_role.id),
                "role_name": matched_role.name,
                "matched": True,
                "match_type": "AI_MATCHED_NAME",
            }

    singular_candidate = normalized_name.rstrip(
        "s"
    )

    for role_name, role in role_indexes[
        "by_name"
    ].items():
        if role_name.rstrip("s") == singular_candidate:
            return {
                "role_id": str(role.id),
                "role_name": role.name,
                "matched": True,
                "match_type": "NORMALIZED_NAME",
            }

    return {
        "role_id": None,
        "role_name": str(role_value).strip(),
        "matched": False,
        "match_type": "ROLE_NOT_FOUND",
    }


def _normalize_amount(
    value,
    *,
    is_unlimited,
    is_allowed,
):
    if not is_allowed:
        return Decimal("0.00")

    if is_unlimited:
        return None

    if value is None:
        return None

    try:
        amount = Decimal(
            str(value)
            .replace(",", "")
            .strip()
        )
    except (
        InvalidOperation,
        TypeError,
        ValueError,
    ):
        return None

    if amount < 0:
        return None

    return amount


# ============================================================
# Backend duplicate/conflict detection
# ============================================================

def _rule_group_key(rule):
    normalized_conditions = tuple(
        sorted(
            str(condition).strip().casefold()
            for condition in rule.get(
                "conditions",
                []
            )
            if condition
        )
    )

    return (
        rule.get("resolved_role_id")
        or "__DEFAULT_EMPLOYEE__",
        rule.get("category"),
        rule.get("currency"),
        rule.get("limit_period"),
        rule.get("domestic_or_international"),
        normalized_conditions,
    )


def _detect_duplicates_and_conflicts(rules):
    grouped_rules = {}

    for rule in rules:
        grouped_rules.setdefault(
            _rule_group_key(rule),
            [],
        ).append(rule)

    warnings = []
    conflicts = []

    for _, group in grouped_rules.items():
        if len(group) < 2:
            continue

        values = {
            (
                rule.get("max_amount"),
                rule.get("is_unlimited"),
                rule.get("is_allowed"),
            )
            for rule in group
        }

        if len(values) == 1:
            for duplicate in group[1:]:
                duplicate["duplicate_rule"] = True
                duplicate["review_required"] = True

                existing_reason = duplicate.get(
                    "review_reason"
                )

                duplicate["review_reason"] = (
                    f"{existing_reason} | Probable duplicate rule."
                    if existing_reason
                    else "Probable duplicate rule."
                )

                warnings.append({
                    "code": "DUPLICATE_RULE",
                    "message": (
                        "A probable duplicate policy rule "
                        "was extracted."
                    ),
                    "rule_index": duplicate["rule_index"],
                    "source_text": duplicate.get(
                        "source_text"
                    ),
                })

        else:
            conflict = {
                "role": (
                    group[0].get("resolved_role_name")
                    or group[0].get("role")
                ),
                "category": group[0].get("category"),
                "values": [
                    (
                        "Unlimited"
                        if rule.get("is_unlimited")
                        else (
                            "Prohibited"
                            if not rule.get("is_allowed")
                            else (
                                f"{rule.get('max_amount')} "
                                f"{rule.get('currency') or ''}"
                            ).strip()
                        )
                    )
                    for rule in group
                ],
                "message": (
                    "Conflicting limits were found for the same "
                    "role, category and conditions."
                ),
                "source_texts": [
                    rule.get("source_text")
                    for rule in group
                    if rule.get("source_text")
                ],
                "rule_indexes": [
                    rule["rule_index"]
                    for rule in group
                ],
            }

            conflicts.append(conflict)

            for rule in group:
                rule["review_required"] = True
                rule["review_reason"] = (
                    "Conflicting policy rules require review."
                )

    return warnings, conflicts


# ============================================================
# Gemini execution
# ============================================================

def _parse_response(response):
    """
    Parse and validate Gemini structured output.
    """

    parsed = getattr(
        response,
        "parsed",
        None,
    )

    if parsed is not None:

        if isinstance(
            parsed,
            ExtractedPolicyDocument,
        ):
            return parsed

        try:
            return ExtractedPolicyDocument.model_validate(
                parsed
            )

        except ValidationError as exc:
            error_details = json.dumps(
                exc.errors(
                    include_url=False,
                    include_context=False,
                ),
                indent=2,
                ensure_ascii=False,
                default=str,
            )

            raise ValueError(
                "Gemini parsed output did not match "
                "the policy schema.\n\n"
                f"{error_details}"
            ) from exc

    response_text = getattr(
        response,
        "text",
        None,
    )

    if not response_text:
        raise ValueError(
            "Gemini returned an empty response."
        )

    try:
        return (
            ExtractedPolicyDocument
            .model_validate_json(
                response_text
            )
        )

    except ValidationError as exc:
        error_details = json.dumps(
            exc.errors(
                include_url=False,
                include_context=False,
            ),
            indent=2,
            ensure_ascii=False,
            default=str,
        )

        raise ValueError(
            "Gemini returned JSON that did not match "
            "the policy schema.\n\n"
            f"{error_details}"
        ) from exc


def _structured_config():
    return types.GenerateContentConfig(
        temperature=0.1,
        response_mime_type="application/json",
        response_schema=ExtractedPolicyDocument,
        max_output_tokens=65536,
    )


def _process_text_document(
    *,
    client,
    prompt,
    document_text,
):
    response = client.models.generate_content(
        model=_get_policy_model(),
        contents=[
            prompt,
            (
                "\n\nPOLICY DOCUMENT CONTENT:\n\n"
                f"{document_text}"
            ),
        ],
        config=_structured_config(),
    )

    return _parse_response(response)


def _process_binary_document(
    *,
    client,
    prompt,
    file_path,
    mime_type,
):
    uploaded = client.files.upload(
        file=file_path,
        config={
            "mime_type": mime_type,
        },
    )

    try:
        response = client.models.generate_content(
            model=_get_policy_model(),
            contents=[
                uploaded,
                prompt,
            ],
            config=_structured_config(),
        )

        return _parse_response(response)

    finally:
        try:
            client.files.delete(
                name=uploaded.name
            )
        except Exception:
            pass


# ============================================================
# Preview normalization
# ============================================================

def _normalize_document(
    *,
    extracted_document,
    company,
):
    currency_indexes = _get_currency_indexes()
    role_indexes = _get_role_indexes(company)

    base_currency = _get_company_base_currency(
        company
    )

    resolved_policy_currency, policy_currency_source = (
        _resolve_currency(
            extracted_document.policy_currency,
            currency_indexes=currency_indexes,
            base_currency=base_currency,
        )
    )

    warnings = [
        warning.model_dump()
        for warning in extracted_document.warnings
    ]

    normalized_rules = []

    for index, extracted_rule in enumerate(
        extracted_document.rules
    ):
        category, unknown_category = (
            _normalize_category(
                extracted_rule.category,
                extracted_rule.original_category,
            )
        )

        role_result = _resolve_role(
            role_value=extracted_rule.role,
            ai_role_match=extracted_rule.role_match,
            role_indexes=role_indexes,
        )

        currency_result, currency_source = (
            _resolve_currency(
                extracted_rule.currency,
                currency_indexes=currency_indexes,
                policy_currency=(
                    resolved_policy_currency.code
                    if resolved_policy_currency
                    else None
                ),
                base_currency=base_currency,
            )
        )

        amount = _normalize_amount(
            extracted_rule.max_amount,
            is_unlimited=extracted_rule.is_unlimited,
            is_allowed=extracted_rule.is_allowed,
        )

        review_required = (
            extracted_rule.review_required
        )

        review_reasons = []

        if extracted_rule.review_reason:
            review_reasons.append(
                extracted_rule.review_reason
            )

        if unknown_category:
            review_required = True
            review_reasons.append(
                "Category could not be matched confidently."
            )

            warnings.append({
                "code": "UNKNOWN_CATEGORY",
                "message": (
                    f"Rule {index + 1} was mapped "
                    "to miscellaneous."
                ),
                "rule_index": index,
                "source_text": extracted_rule.source_text,
            })

        if not role_result["matched"]:
            review_required = True
            review_reasons.append(
                "The named role was not found in the company."
            )

            warnings.append({
                "code": "ROLE_NOT_FOUND",
                "message": (
                    f"Role '{extracted_rule.role}' does not "
                    "exist as an active company role."
                ),
                "rule_index": index,
                "source_text": extracted_rule.source_text,
            })

        if currency_result is None:
            review_required = True
            review_reasons.append(
                "Currency could not be resolved."
            )

            warnings.append({
                "code": "UNCLEAR_CURRENCY",
                "message": (
                    f"Currency could not be resolved "
                    f"for rule {index + 1}."
                ),
                "rule_index": index,
                "source_text": extracted_rule.source_text,
            })

        if (
            extracted_rule.is_allowed
            and not extracted_rule.is_unlimited
            and amount is None
        ):
            review_required = True
            review_reasons.append(
                "Maximum amount could not be determined."
            )

            warnings.append({
                "code": "UNCLEAR_AMOUNT",
                "message": (
                    f"Amount could not be determined "
                    f"for rule {index + 1}."
                ),
                "rule_index": index,
                "source_text": extracted_rule.source_text,
            })

        if extracted_rule.confidence < 0.75:
            review_required = True
            review_reasons.append(
                "Extraction confidence is below 0.75."
            )

        if extracted_rule.translation_confidence < 0.75:
            review_required = True
            review_reasons.append(
                "Translation confidence is below 0.75."
            )

            warnings.append({
                "code": "UNCLEAR_TRANSLATION",
                "message": (
                    f"Translation requires review "
                    f"for rule {index + 1}."
                ),
                "rule_index": index,
                "source_text": extracted_rule.source_text,
            })

        normalized_rules.append({
            "rule_index": index,

            "role": extracted_rule.role,
            "role_original_text": (
                extracted_rule.role_original_text
            ),
            "resolved_role_id": role_result["role_id"],
            "resolved_role_name": role_result["role_name"],
            "role_matched": role_result["matched"],
            "role_match_type": role_result["match_type"],

            "category": category,
            "original_category": (
                extracted_rule.original_category
                or extracted_rule.category
            ),

            "max_amount": (
                format(amount, "f")
                if amount is not None
                else None
            ),

            "currency": (
                currency_result.code
                if currency_result
                else None
            ),
            "currency_name": (
                currency_result.name
                if currency_result
                else None
            ),
            "currency_symbol": (
                currency_result.symbol
                if currency_result
                else None
            ),
            "currency_country": (
                currency_result.country
                if currency_result
                else None
            ),
            "currency_resolution_source": currency_source,

            "is_unlimited": extracted_rule.is_unlimited,
            "is_allowed": extracted_rule.is_allowed,

            "description": extracted_rule.description,
            "description_original_language": (
                extracted_rule
                .description_original_language
            ),

            "reason": extracted_rule.reason,
            "reason_original_language": (
                extracted_rule.reason_original_language
            ),

            "source_text": extracted_rule.source_text,
            "translated_source_text": (
                extracted_rule.translated_source_text
            ),
            "source_language": extracted_rule.source_language,
            "source_language_code": (
                extracted_rule.source_language_code
            ),

            "limit_period": extracted_rule.limit_period,

            "conditions": extracted_rule.conditions,
            "conditions_original_language": (
                extracted_rule
                .conditions_original_language
            ),

            "required_documents": (
                extracted_rule.required_documents
            ),
            "required_documents_original_language": (
                extracted_rule
                .required_documents_original_language
            ),

            "receipt_required": (
                extracted_rule.receipt_required
            ),
            "approval_required": (
                extracted_rule.approval_required
            ),

            "domestic_or_international": (
                extracted_rule
                .domestic_or_international
            ),

            "country_scope": extracted_rule.country_scope,
            "department_scope": (
                extracted_rule.department_scope
            ),
            "employee_grade_scope": (
                extracted_rule.employee_grade_scope
            ),

            "travel_class": extracted_rule.travel_class,
            "distance_rate": extracted_rule.distance_rate,
            "distance_unit": extracted_rule.distance_unit,
            "tax_included": extracted_rule.tax_included,

            "confidence": float(
                extracted_rule.confidence
            ),
            "translation_confidence": float(
                extracted_rule.translation_confidence
            ),

            "review_required": review_required,
            "review_reason": (
                " | ".join(
                    dict.fromkeys(review_reasons)
                )
                if review_reasons
                else None
            ),

            "duplicate_rule": False,
        })

    backend_warnings, backend_conflicts = (
        _detect_duplicates_and_conflicts(
            normalized_rules
        )
    )

    warnings.extend(backend_warnings)

    conflicts = [
        conflict.model_dump()
        for conflict in extracted_document.conflicts
    ]

    conflicts.extend(backend_conflicts)

    return {
        "document_language": (
            extracted_document.document_language
        ),
        "document_language_code": (
            extracted_document.document_language_code
        ),
        "contains_multiple_languages": (
            extracted_document.contains_multiple_languages
        ),
        "languages_detected": [
            item.model_dump()
            for item in extracted_document.languages_detected
        ],

        "policy_name": extracted_document.policy_name,
        "policy_version": extracted_document.policy_version,
        "effective_date": extracted_document.effective_date,
        "expiry_date": extracted_document.expiry_date,

        "policy_currency": (
            resolved_policy_currency.code
            if resolved_policy_currency
            else None
        ),
        "policy_currency_name": (
            resolved_policy_currency.name
            if resolved_policy_currency
            else None
        ),
        "policy_currency_source": policy_currency_source,

        "company_base_currency": (
            base_currency.code
            if base_currency
            else None
        ),

        "document_summary": (
            extracted_document.document_summary
        ),
        "document_summary_original_language": (
            extracted_document
            .document_summary_original_language
        ),

        "document_level_conditions": (
            extracted_document.document_level_conditions
        ),
        "document_level_conditions_original_language": (
            extracted_document
            .document_level_conditions_original_language
        ),

        "rules": normalized_rules,
        "rules_found": len(normalized_rules),

        "rules_requiring_review": sum(
            1
            for rule in normalized_rules
            if rule["review_required"]
        ),

        "warnings": warnings,
        "warning_count": len(warnings),

        "conflicts": conflicts,
        "conflict_count": len(conflicts),
    }


# ============================================================
# Public service
# ============================================================

def extract_policy_document(
    *,
    uploaded_file,
    company,
):
    """
    Extract and normalize a policy document.

    This function:
    - does not create PolicyCategoryRule rows;
    - does not create PolicyDocumentImport;
    - returns preview JSON for later review and storage.

    PolicyDocumentImport will be created by the upload API.
    """

    extension = _validate_uploaded_file(
        uploaded_file
    )

    temporary_path = None

    try:

        temporary_path = _save_temporary_file(
            uploaded_file,
            extension,
        )

        prompt = build_policy_document_prompt(
            company
        )

        client = _get_gemini_client()

        # -----------------------------------------
        # Process document
        # -----------------------------------------

        if extension == ".docx":

            text = _extract_docx_text(
                temporary_path
            )

            extracted_document = _process_text_document(
                client=client,
                prompt=prompt,
                document_text=text,
            )

        elif extension == ".txt":

            text = _extract_txt_text(
                temporary_path
            )

            extracted_document = _process_text_document(
                client=client,
                prompt=prompt,
                document_text=text,
            )

        else:

            mime_type = (
                getattr(uploaded_file, "content_type", None)
                or mimetypes.guess_type(
                    getattr(
                        uploaded_file,
                        "name",
                        "",
                    )
                )[0]
            )

            if not mime_type:
                raise ValueError(
                    "Unable to determine document MIME type."
                )

            extracted_document = _process_binary_document(
                client=client,
                prompt=prompt,
                file_path=temporary_path,
                mime_type=mime_type,
            )

        # -----------------------------------------
        # Normalize AI Output
        # -----------------------------------------

        preview = _normalize_document(
            extracted_document=extracted_document,
            company=company,
        )

        # -----------------------------------------
        # Company Language Settings
        # -----------------------------------------

        language_settings = (
            get_company_output_language(
                company
            )
        )

        # -----------------------------------------
        # Basic Metadata
        # -----------------------------------------

        preview["original_filename"] = (
            getattr(
                uploaded_file,
                "name",
                "",
            )
        )

        preview["ai_model"] = (
            _get_policy_model()
        )

        preview["prompt_version"] = "v2.0"

        preview["document_type"] = (
            preview.get(
                "document_type",
                "Expense Policy",
            )
        )

        # -----------------------------------------
        # Language
        # -----------------------------------------

        preview["output_language"] = (
            language_settings["name"]
        )

        preview["output_language_code"] = (
            language_settings["code"]
        )

        preview["preserve_original_text"] = (
            language_settings[
                "preserve_original_text"
            ]
        )

        # -----------------------------------------
        # AI Metadata
        # -----------------------------------------

        preview["translation"] = {

            "source_language": preview.get(
                "document_language"
            ),

            "source_language_code": preview.get(
                "document_language_code"
            ),

            "target_language": (
                language_settings["name"]
            ),

            "target_language_code": (
                language_settings["code"]
            ),

        }

        # Gemini will populate this later when we
        # enhance the prompt.
        preview["document_quality"] = {

            "score": None,

            "status": "NOT_ASSESSED",

            "issues": [],

        }

        # -----------------------------------------
        # Store Complete AI JSON
        # -----------------------------------------

        preview["ai_json"] = (
            extracted_document.model_dump(
                mode="json"
            )
        )

        return {

            "success": True,

            "preview": preview,

        }

    except Exception as exc:

        return {

            "success": False,

            "error": str(exc),

        }

    finally:

        if (
            temporary_path
            and os.path.exists(
                temporary_path
            )
        ):

            try:

                os.remove(
                    temporary_path
                )

            except OSError:

                pass